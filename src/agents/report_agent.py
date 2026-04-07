"""
Report Agent — Phase 4.

배터리 시장 전략 분석 최종 보고서를 생성하고 Word(.docx)로 저장한다.
노트북(Battery-Comparative-SWOT-ReportAgent.py)의 report_agent와
publish_report 로직을 모듈화한 버전이다.

입력:
  - user_request       : 사용자 요청 (강조 포인트)
  - market_context     : MarketContext (Market Agent 출력)
  - skon_result        : StrategyAgentOutput
  - catl_result        : StrategyAgentOutput
  - comparative_swot   : ComparativeSWOTOutput dict
  - human_feedback     : HITL #3 거절 피드백 (재실행 시)
  - final_revision_mode: 최대 리뷰 횟수 도달 시 최종 수정 모드 플래그

출력:
  FinalReportOutput (dict)

Word 저장:
  build_word_report(report_draft, comparative_swot, output_path)
  ※ python-docx 패키지 필요: pip install python-docx
"""

from __future__ import annotations
import json
from pathlib import Path
from typing import Dict, List

from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field

from rag.source_metadata import resolve_source_metadata
from schemas.agent_io import SourceRecord, StrategyAgentOutput
from schemas.market_context import MarketContext
from logging_utils import get_logger


logger = get_logger("report_agent")


# ─────────────────────────────────────────────
# 출력 스키마
# ─────────────────────────────────────────────

class CompanySection(BaseModel):
    company: str = Field(description="회사명")
    portfolio_diversification: List[str] = Field(description="포트폴리오 다각화 내용")
    core_competencies: List[str] = Field(description="핵심 경쟁력")
    strategic_direction: str = Field(description="현재 전략 방향")
    key_watchpoints: List[str] = Field(description="추가로 볼 포인트")


class FinalReportOutput(BaseModel):
    title: str = Field(description="보고서 제목")
    summary: str = Field(
        description="기업 자체의 절대 평가보다 시장 변화 대응 적합성에 초점을 둔 핵심 메시지 요약. 반 페이지를 넘지 않는 분량"
    )
    market_background: List[str] = Field(
        description="Market Agent가 포착한 배터리 시장 환경 변화와 전략적 의미. 각 bullet은 수치, 연도, 지역, 규제, 전략적 해석을 포함한 상세 문장으로 작성"
    )
    sk_on_section: CompanySection = Field(
        description="SK On이 시장 변화에 어떻게 대응하고 있는지 보여주는 포트폴리오 다각화 및 핵심 경쟁력"
    )
    catl_section: CompanySection = Field(
        description="CATL이 시장 변화에 어떻게 대응하고 있는지 보여주는 포트폴리오 다각화 및 핵심 경쟁력"
    )
    comparative_swot_focus_points: List[str] = Field(
        description="4-1에 들어갈 SWOT 비교 기준 설명. 각 bullet은 무엇을 비교하는 기준인지와 왜 중요한지만 설명하고, 특정 기업 우위 판단은 쓰지 않음"
    )
    comparative_swot_company_comparison: List[str] = Field(
        description="4-2에 들어갈 S/W/O/T 별 기업 비교. 4-1의 기준을 바탕으로 SK On과 CATL의 대응 차이를 비교"
    )
    integrated_implications: List[str] = Field(
        description="현재 우위 기업, 우위가 뒤집힐 조건, SK On 우선 과제, CATL 리스크 모니터링 포인트, 의사결정자가 볼 핵심 신호를 포함한 종합 시사점"
    )
    references: List[str] = Field(description="실제 활용 자료 목록. 지정 형식 준수")


_REPORT_PROMPT = """
당신은 의사결정자용 배터리 시장 전략 보고서 작성자입니다.
아래 입력만 사용해 보고서 초안을 작성하세요.

보고서 목차:
1. Summary
2. 시장 배경(배터리 시장 환경 변화)
3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력
   - SK ON
   - CATL
4. Comparative SWOT
   - SWOT 비교 기준
   - SWOT 기업별 비교
   - SWOT 비교 요약 table
5. 종합 시사점
6. Reference

작성 규칙:
- Summary는 보고서 전체 핵심 메시지 중심으로 작성하고, 1/2 페이지를 넘지 않게 압축할 것
- Summary만 임원용 요약으로 짧게 작성하고, 나머지 본문 섹션은 절대 과도하게 요약하지 말 것
- summary를 포함한 모든 서술형 문장, bullet, 표 설명은 반드시 한국어로 작성할 것
- 모든 리스트 항목은 명사형 메모나 짧은 꼬리표가 아니라, 주어와 서술어가 있는 완결된 한국어 문장으로 작성할 것
- 각 bullet은 가능하면 2~3문장으로 구성해 근거와 해석이 함께 드러나게 할 것
- "포트폴리오 다각화", "핵심 경쟁력", "SWOT 비교 기준", "SWOT 기업별 비교", "종합 시사점"의 각 항목은 한 줄짜리 키워드 나열로 쓰지 말 것
- strategic_direction은 짧은 구가 아니라 3~5문장 분량의 연결된 설명문으로 작성할 것
- key_watchpoints도 단문 키워드가 아니라, 왜 주의해서 봐야 하는지 설명이 들어간 완결 문장으로 작성할 것
- 회사명, 배터리 화학계열(LFP, NCM), 제도명(IRA), 단위, 고유명사 외에는 영어 문장을 쓰지 말 것
- user_request는 기본 목차를 바꾸는 명령이 아니라, 이번 보고서에서 더 강조해야 할 관점을 알려주는 입력이다
- 예를 들어 규제, 점유율, 원가, 기술, ESS, 캐즘 같은 키워드가 있으면 해당 관점을 본문과 종합 시사점에서 더 선명하게 드러낼 것
- 보고서 전체 논리는 "기업 자체가 얼마나 좋은가"보다 "Market Agent가 포착한 시장 변화에 얼마나 잘 맞게 대응하고 있는가" 중심으로 전개할 것

섹션별 품질 기준 (반드시 준수):

[시장 배경]
- 각 bullet은 반드시 수치(%, GWh, $/kWh, 성장률 등) + 해당 수치가 배터리 기업에 갖는 전략적 의미를 함께 포함할 것
- "EV 성장이 둔화됐다"처럼 사실 나열로 끝내지 말고, "이것이 한국 배터리 3사에 미치는 구체적 영향"까지 이어서 쓸 것
- market_context의 detailed_analysis, key_narrative를 최대한 활용해 실제 발견한 정보를 풀어쓸 것

[기업별 섹션]
- portfolio_diversification: 각 항목에 규모(GWh·억 원), 시기(연도), 지역·고객 중 2개 이상 포함
- core_competencies: "기술력이 뛰어나다" 식의 서술 금지. 반드시 수치 또는 비교 대상을 명시
- strategic_direction: 현재 전략과 시장 트렌드 간 정합(aligned)/불일치(misaligned) 판단 포함
- key_watchpoints: 각 항목은 "무엇을 언제까지 주시해야 하는지"와 "그 결과가 어떤 전략 판단을 바꾸는지"가 드러나는 문장으로 작성

[Comparative SWOT 4-2]
- 각 비교 항목은 comparative_swot의 reversal_condition을 반드시 반영할 것
- A기업의 강점이 B기업에 구체적으로 어떤 압력을 주는지 메커니즘이 보이게 작성

[종합 시사점]
- 각 항목은 "관찰 가능한 시장 신호 → 그때 취해야 할 전략 행동" 형식으로 작성
- ✗ 금지: "SK On은 원가 경쟁력을 강화해야 한다"
- ✓ 필수: "배터리 팩 가격이 $100/kWh 이하로 정착되면 NCM 중심 라인업을 LFP 비중 X% 이상으로 전환하지 않는 한 수익성 방어 불가 → 2025년 내 LFP 상용화 로드맵 공시 여부가 핵심 체크포인트"
- 최소 5개, 아래 5가지를 각각 포함:
  ① 현재 기준 우위 기업 + 수치 근거
  ② 역전 조건 + 관찰 가능한 임계값
  ③ SK On 최우선 대응 과제 + 기한 또는 조건
  ④ CATL 핵심 리스크 + 모니터링 지표
  ⑤ 의사결정자가 주시해야 할 조기경보 시장 신호

- 기업별 섹션은 MarketContext를 기준으로 각 기업의 포트폴리오 다각화와 핵심 경쟁력이 시장 변화에 얼마나 맞는지 쓰고, 포트폴리오 다각화, 핵심 경쟁력, 전략 방향, watchpoint를 분리해서 작성할 것
- SK On과 CATL의 각 섹션은 입력으로 주어진 strategy_result의 content 원문을 우선 사용하고, content에 있는 수치, 설비, 고객, 지역, JV, 기술, 정책 대응을 최대한 보존할 것
- 보고서 본문에서 원문 evidence를 지나치게 축약하지 말고, 필요한 경우 한 bullet 안에서 2~3문장으로 상세히 설명할 것
- 각 리스트 필드는 최소 4개 이상 작성하려고 노력할 것. 정보가 충분하면 6개 이상 작성해도 된다
- Comparative SWOT에서는 swot_focus_points, strategic_interactions, swot_comparison_table을 모두 반영할 것
- 4-1은 비교 기준 설명 전용이다. 무엇을 어떤 기준으로 비교하는지, 왜 중요한지만 적고 특정 기업 우위나 비교 결과는 쓰지 말 것
- 4-2는 4-1의 기준을 바탕으로 S/W/O/T 별 기업 비교를 적을 것
- A기업의 강점이 B기업의 위협이나 약점으로 직결되는 전략적 상호작용이 드러나야 함
- references는 제공된 형식화된 목록만 사용하고, 실제 활용한 자료만 남길 것
- references는 아래 Allowed references 목록의 문장을 그대로 사용하고, 문구를 임의로 바꾸거나 축약하지 말 것
- review_2_feedback가 있으면 해당 보완 지시를 반영할 것
- human_feedback가 있으면 반드시 반영할 것
- final_revision_mode가 true이면 최대한 완성도 높게 작성할 것
- 아래 raw evidence 입력은 보고서 본문을 풍부하게 쓰기 위한 근거 묶음이므로 적극 활용할 것
- 반드시 JSON으로만 답할 것

Market context:
{market_context}

Market raw evidence:
{market_evidence}

Comparative SWOT:
{comparative_swot}

SK On 전략 분석:
{skon_result}

SK On raw evidence:
{skon_raw_evidence}

CATL 전략 분석:
{catl_result}

CATL raw evidence:
{catl_raw_evidence}

Allowed references:
{references}

user_request:
{user_request}

review2_feedback_from_human_review_2:
{review_2_feedback}

human_feedback_from_previous_review:
{human_feedback}

final_revision_mode:
{final_revision_mode}

{format_instructions}
"""


# ─────────────────────────────────────────────
# 참고문헌 빌더
# ─────────────────────────────────────────────

_MARKET_SECTION_KEYS = [
    "ev_growth_slowdown", "market_share_ranking", "lfp_ncm_trend",
    "ess_hev_growth", "regulatory_status", "cost_competitiveness",
]
_STRATEGY_AXIS_KEYS = [
    "ev_response", "market_position", "tech_portfolio",
    "ess_strategy", "regulatory_risk", "cost_structure",
]


def _collect_used_source_ids(
    market_context: MarketContext,
    skon_result: StrategyAgentOutput,
    catl_result: StrategyAgentOutput,
) -> set[str]:
    """각 분석 섹션의 source_ids를 수집 — 실제 콘텐츠 생성에 사용된 출처 ID 집합."""
    used: set[str] = set()

    for key in _MARKET_SECTION_KEYS:
        section = market_context.get(key)
        if isinstance(section, dict):
            for sid in section.get("source_ids", []):
                used.add(str(sid).strip())

    for result in (skon_result, catl_result):
        for key in _STRATEGY_AXIS_KEYS:
            axis = result.get(key)
            if isinstance(axis, dict):
                for sid in axis.get("source_ids", []):
                    used.add(str(sid).strip())

    return used


def _build_references(
    market_context: MarketContext,
    skon_result: StrategyAgentOutput,
    catl_result: StrategyAgentOutput,
    collected_sources: list[dict] | None = None,
) -> List[str]:
    """섹션 source_ids에 실제로 등장한 소스만 참고문헌으로 포함."""
    used_ids = _collect_used_source_ids(market_context, skon_result, catl_result)

    # source_id → SourceRecord 역인덱스
    source_map: dict[str, dict] = {}
    if collected_sources:
        # Primary: LangGraph state에서 누적된 소스 풀
        for src in collected_sources:
            sid = str(src.get("source_id", "")).strip()
            if sid:
                source_map.setdefault(sid, src)
    else:
        # Fallback: 중첩 dict에서 직접 추출 (state 미사용 시)
        for src in market_context.get("source_records", []):
            sid = str(src.get("source_id", "")).strip()
            if sid:
                source_map.setdefault(sid, src)
        for src in list(skon_result.get("sources", [])) + list(catl_result.get("sources", [])):
            sid = str(src.get("source_id", "")).strip()
            if sid:
                source_map.setdefault(sid, src)

    seen: set[str] = set()
    refs: list[str] = []

    # 실제 사용된 source_id만 포함
    for sid in sorted(used_ids):
        src = source_map.get(sid)
        if src:
            _append_source(src, seen, refs)
        else:
            # source_map에 없으면 ID 자체로 resolve 시도 (PDF 청크 등)
            resolved = resolve_source_metadata(source_id=sid)
            text = resolved["reference_text"].strip()
            if text and text not in seen:
                seen.add(text)
                refs.append(text)

    # market_context["references"] 중 source_id 기반 매칭이 없는 항목도 보완
    for item in market_context.get("references", []):
        sid = str(item.get("source_id", "")).strip()
        if sid and sid not in used_ids:
            continue  # 실제 사용 안 된 출처는 제외
        formatted = str(item.get("formatted_reference", "")).strip()
        if formatted and formatted not in seen:
            seen.add(formatted)
            refs.append(formatted)

    return refs or ["내부 리서치 문서 및 공개 출처 참조"]


def _append_source(src: dict, seen: set[str], refs: list[str]) -> None:
    """SourceRecord 하나를 formatted reference 문자열로 변환 후 중복 없이 추가."""
    source_id = str(src.get("source_id", "")).strip()
    title = str(src.get("title", "")).strip()
    url = str(src.get("url", "")).strip()
    retrieved = str(src.get("retrieved_at", "")).strip()[:10] or "n.d."
    resolved = resolve_source_metadata(source_id=source_id, title=title)
    formatted = resolved["reference_text"].strip()
    if not formatted:
        display = title or resolved["title"] or "Source"
        if str(src.get("source_type", "")) == "web":
            formatted = f"{display}({retrieved}). *{display}*. Web, {url or resolved['url']}"
        else:
            formatted = f"{display}. {url or resolved['url']}".strip(" .")
    if formatted and formatted not in seen:
        seen.add(formatted)
        refs.append(formatted)


def _build_market_evidence(market_context: MarketContext) -> str:
    """Build a dense market evidence bundle for the report agent."""

    lines: list[str] = []
    for key in [
        "ev_growth_slowdown",
        "market_share_ranking",
        "lfp_ncm_trend",
        "ess_hev_growth",
        "regulatory_status",
        "cost_competitiveness",
    ]:
        section = market_context.get(key, {})
        if not isinstance(section, dict):
            continue
        lines.append(f"[{key}]")
        key_narrative = section.get("key_narrative")
        if key_narrative:
            lines.append(f"key_narrative: {key_narrative}")
        detailed_analysis = section.get("detailed_analysis")
        if detailed_analysis:
            lines.append(f"detailed_analysis: {detailed_analysis}")
        lines.append(json.dumps(section, ensure_ascii=False, indent=2))
    return "\n".join(lines)


def _build_company_evidence(result: StrategyAgentOutput) -> str:
    """Build a dense company evidence bundle from raw strategy outputs."""

    lines: list[str] = []
    for key in [
        "ev_response",
        "market_position",
        "tech_portfolio",
        "ess_strategy",
        "regulatory_risk",
        "cost_structure",
    ]:
        section = result.get(key, {})
        if not isinstance(section, dict):
            continue
        lines.append(f"[{key}]")
        content = str(section.get("content", "")).strip()
        if content:
            lines.append(content)
        source_ids = section.get("source_ids", [])
        if source_ids:
            lines.append(f"source_ids: {source_ids}")
    return "\n\n".join(lines)


# ─────────────────────────────────────────────
# 실행 함수
# ─────────────────────────────────────────────

async def run_report_agent(
    user_request: str,
    market_context: MarketContext,
    skon_result: StrategyAgentOutput,
    catl_result: StrategyAgentOutput,
    comparative_swot: dict,
    review_2_feedback: str = "",
    human_feedback: str = "",
    final_revision_mode: bool = False,
    collected_sources: list[dict] | None = None,
) -> dict:
    """
    Report Agent 실행. 보고서 초안(JSON)을 생성한다.

    Args:
        user_request        : 사용자 요청(강조 포인트)
        market_context      : Market Agent 출력
        skon_result         : SKON Strategy Agent 출력
        catl_result         : CATL Strategy Agent 출력
        comparative_swot    : Comparative SWOT Agent 출력
        review_2_feedback   : HITL #2 재조사/보완 지시
        human_feedback      : HITL #3 거절 피드백 (재실행 시)
        final_revision_mode : 리뷰 한계 도달 시 최종 수정 모드

    Returns:
        FinalReportOutput을 dict로 변환한 결과
    """
    with logger.node_span("report_agent", {
        "has_feedback": bool(human_feedback),
        "final_revision": final_revision_mode,
    }):
        references = _build_references(
            market_context, skon_result, catl_result,
            collected_sources=collected_sources,
        )
        market_evidence = _build_market_evidence(market_context)
        skon_raw_evidence = _build_company_evidence(skon_result)
        catl_raw_evidence = _build_company_evidence(catl_result)

        parser = JsonOutputParser(pydantic_object=FinalReportOutput)
        prompt = PromptTemplate(
            template=_REPORT_PROMPT,
            input_variables=[
                "market_context", "comparative_swot",
                "skon_result", "catl_result",
                "references", "human_feedback", "final_revision_mode",
                "market_evidence", "skon_raw_evidence", "catl_raw_evidence", "user_request",
                "review_2_feedback",
            ],
            partial_variables={"format_instructions": parser.get_format_instructions()},
        )
        llm   = ChatOpenAI(model="gpt-4o", temperature=0)
        chain = prompt | llm | parser

        result = await chain.ainvoke({
            "market_context":      json.dumps(market_context,    ensure_ascii=False, indent=2),
            "comparative_swot":    json.dumps(comparative_swot,  ensure_ascii=False, indent=2),
            "skon_result":         json.dumps(skon_result,        ensure_ascii=False, indent=2),
            "catl_result":         json.dumps(catl_result,        ensure_ascii=False, indent=2),
            "references":          json.dumps(references,         ensure_ascii=False, indent=2),
            "market_evidence":     market_evidence,
            "skon_raw_evidence":   skon_raw_evidence,
            "catl_raw_evidence":   catl_raw_evidence,
            "user_request":        user_request,
            "review_2_feedback":   review_2_feedback or "No review_2 feedback.",
            "human_feedback":      human_feedback or "No prior human feedback.",
            "final_revision_mode": final_revision_mode,
        })
        normalized = result if isinstance(result, dict) else result.model_dump()
        normalized["references"] = references
        return normalized


# ─────────────────────────────────────────────
# Word 문서 저장
# ─────────────────────────────────────────────

def build_word_report(
    report: Dict,
    swot: Dict,
    output_path: Path,
) -> None:
    """
    보고서 JSON + SWOT JSON → Word(.docx) 저장.

    Args:
        report      : FinalReportOutput dict
        swot        : ComparativeSWOTOutput dict
        output_path : 저장 경로 (.docx)

    Raises:
        ImportError : python-docx 미설치 시
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise ImportError(
            "Word 보고서 생성에는 python-docx가 필요합니다: pip install python-docx"
        ) from e

    ACCENT_COLOR       = RGBColor(0xFB, 0x77, 0x62)
    SEPARATOR_COLOR    = "D9D9D9"
    TITLE_SEPARATOR_COLOR = "FB7762"
    TABLE_HEADER_FILL  = "FDE3DE"
    TABLE_FIRST_COL_FILL = "FEF1EE"

    def _remove_numbering(paragraph):
        p_pr  = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)

    def _add_bullet_items(doc: Document, items: List[str]):
        for item in items:
            p = doc.add_paragraph()
            _remove_numbering(p)
            p.add_run(f"· {item}")

    def _add_toc_items(doc: Document, items: List[tuple[str, int]]):
        # 목차는 대주제/소주제를 시각적으로 구분해 사람이 바로 읽기 쉽게 만든다.
        for text, level in items:
            p = doc.add_paragraph()
            _remove_numbering(p)
            if level > 0:
                p.paragraph_format.left_indent = Pt(18 * level)
            p.add_run(f"· {text}")

    def _add_clean_heading(doc: Document, text: str, level: int = 1):
        p = doc.add_paragraph(style=f"Heading {level}")
        _remove_numbering(p)
        run = p.add_run(text)
        run.font.color.rgb = ACCENT_COLOR
        return p

    def _set_run_size(run, size_pt: int):
        run.font.size = Pt(size_pt)

    def _add_title(doc: Document, text: str):
        p = doc.add_paragraph()
        _remove_numbering(p)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = ACCENT_COLOR
        _set_run_size(run, 22)
        return p

    def _add_separator_line(
        doc: Document,
        color: str = SEPARATOR_COLOR,
        val: str = "single",
        size: str = "6",
    ):
        p = doc.add_paragraph()
        _remove_numbering(p)
        p_pr  = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), val)
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _set_cell_background(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _add_company_section(doc: Document, heading: str, section: Dict):
        _add_clean_heading(doc, heading, level=2)
        _add_clean_heading(doc, "포트폴리오 다각화", level=3)
        _add_bullet_items(doc, section.get("portfolio_diversification", []))
        _add_clean_heading(doc, "핵심 경쟁력", level=3)
        _add_bullet_items(doc, section.get("core_competencies", []))
        _add_clean_heading(doc, "전략 방향", level=3)
        strategic_direction = section.get("strategic_direction", "")
        if strategic_direction:
            _add_bullet_items(doc, [strategic_direction])
        watchpoints = section.get("key_watchpoints", [])
        if watchpoints:
            _add_clean_heading(doc, "Watchpoints", level=3)
            _add_bullet_items(doc, watchpoints)

    def _add_swot_table(doc: Document):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "구분"
        hdr[1].text = "A기업 (SK On)"
        hdr[2].text = "B기업 (CATL)"
        hdr[3].text = "전략적 시사점 (비교 평가)"
        for cell in hdr:
            _set_cell_background(cell, TABLE_HEADER_FILL)
        for row in swot.get("swot_comparison_table", []):
            cells = table.add_row().cells
            cells[0].text = row.get("category", "")
            cells[1].text = row.get("company_a_summary", "")
            cells[2].text = row.get("company_b_summary", "")
            cells[3].text = row.get("strategic_implication", "")
            _set_cell_background(cells[0], TABLE_FIRST_COL_FILL)

    # ── 문서 작성 ──────────────────────────────
    doc = Document()

    _add_title(doc, report.get("title", "배터리 시장 전략 분석 보고서"))
    _add_separator_line(doc, color=TITLE_SEPARATOR_COLOR, val="double", size="10")

    _add_clean_heading(doc, "목차", level=1)
    _add_toc_items(doc, [
        ("1. Summary", 0),
        ("2. 시장 배경(배터리 시장 환경 변화)", 0),
        ("3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력", 0),
        ("3-1. SK On", 1),
        ("3-2. CATL", 1),
        ("4. Comparative SWOT", 0),
        ("4-1. SWOT 비교 기준", 1),
        ("4-2. SWOT 기업별 비교", 1),
        ("4-3. SWOT 비교 요약 table", 1),
        ("5. 종합 시사점", 0),
        ("6. Reference", 0),
    ])
    _add_separator_line(doc, color=SEPARATOR_COLOR, val="dashed", size="6")

    _add_clean_heading(doc, "1. Summary", level=1)
    doc.add_paragraph(report.get("summary", ""))

    _add_clean_heading(doc, "2. 시장 배경(배터리 시장 환경 변화)", level=1)
    _add_bullet_items(doc, report.get("market_background", []))

    _add_clean_heading(doc, "3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력", level=1)
    _add_company_section(doc, "3-1. SK On", report.get("sk_on_section", {}))
    _add_company_section(doc, "3-2. CATL",  report.get("catl_section", {}))

    _add_clean_heading(doc, "4. Comparative SWOT", level=1)
    _add_clean_heading(doc, "4-1. SWOT 비교 기준", level=2)
    _add_bullet_items(doc, report.get("comparative_swot_focus_points", []))
    _add_clean_heading(doc, "4-2. SWOT 기업별 비교", level=2)
    _add_bullet_items(doc, report.get("comparative_swot_company_comparison", []))
    _add_clean_heading(doc, "4-3. SWOT 비교 요약 table", level=2)
    _add_swot_table(doc)

    _add_clean_heading(doc, "5. 종합 시사점", level=1)
    _add_bullet_items(doc, report.get("integrated_implications", []))

    _add_clean_heading(doc, "6. Reference", level=1)
    _add_bullet_items(doc, report.get("references", []))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.node_exit("publish_report", duration_sec=0, status="ok",
                     metadata={"path": str(output_path)})
